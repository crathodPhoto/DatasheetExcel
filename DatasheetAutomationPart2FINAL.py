import os
import shutil
import pandas as pd

from datetime import datetime
import win32com.client as win32
from docx import Document
from PIL import Image
from docx.shared import Inches

# -------------------------
# CONFIGURATION - Paths
# -------------------------
destination_folder = r"C:\Users\crathod\Documents\Datasheet Automation\Script Output"
data_package_folder = os.path.join(destination_folder, "Data Package")
excel_template_path = r"C:\Users\crathod\Documents\Datasheet Automation\Datasheet Graph Template 1.xlsm"
word_template_path = r"C:\Users\crathod\Documents\Datasheet Automation\Datasheet Template.docx"

os.makedirs(data_package_folder, exist_ok=True)

# -------------------------
# Load Devices.xlsx
# -------------------------
devices_file = os.path.join(destination_folder, "Devices.xlsx")
devices_df = pd.read_excel(devices_file, sheet_name="Devices")
# Only drop rows where Lot_ID or Dev# are missing (SN and SKU can be empty initially)
devices_df = devices_df.dropna(subset=["Lot_ID", "Dev#"])

# Fill NaN values in SN and SKU with empty strings/default values for processing
devices_df["SN"] = devices_df["SN"].fillna("")
devices_df["SKU"] = devices_df["SKU"].fillna("")

# -------------------------
# Helpers
# -------------------------
def find_device_file(lot_id, dev_num, phrase):
    for filename in os.listdir(os.path.join(destination_folder, "Other")):
        if lot_id in filename and dev_num in filename and phrase in filename:
            return os.path.join(destination_folder, "Other", filename)
    return None

def paste_text_file_fast(sheet, start_row, start_column, lot_id, dev_num, phrase):
    file_path = find_device_file(lot_id, dev_num, phrase)
    if not file_path:
        print(f"File not found for {phrase} - skipping.")
        return None

    with open(file_path, 'r') as f:
        lines = [line.strip() for line in f.readlines()]

    data = [line.split() for line in lines]
    num_rows = len(data)
    num_cols = max(len(row) for row in data)

    for row in data:
        while len(row) < num_cols:
            row.append("")

    start_cell = sheet.Cells(start_row, start_column)
    end_cell = sheet.Cells(start_row + num_rows - 1, start_column + num_cols - 1)
    sheet.Range(start_cell, end_cell).Value = data

    print(f"Fast-pasted {len(data)} rows for {phrase} starting at row {start_row}, column {start_column}")
    return file_path

def clear_old_data(sheet):
    """Clear specific rows before pasting new data."""
    print("Clearing old data in rows 69-72 and 40-43...")
    sheet.Range("A69:CB72").ClearContents()
    sheet.Range("A40:CB43").ClearContents()
    print("Old data cleared.\n")

def resize_image(input_path, output_path, scale_percent):
    with Image.open(input_path) as img:
        new_width = int(img.width * (scale_percent / 100))
        new_height = int(img.height * (scale_percent / 100))
        img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        img.save(output_path)

def update_chart_axes(sheet, chart, chart_number):
    if chart_number == 1:
        axes_config = {
            'primary_y': (sheet.Cells(13, 4).Value, sheet.Cells(13, 5).Value),
            'primary_x': (sheet.Cells(7, 4).Value, sheet.Cells(7, 5).Value),
            'secondary_y': (sheet.Cells(10, 4).Value, sheet.Cells(10, 5).Value)
        }
    elif chart_number == 2:
        axes_config = {
            'primary_y': (sheet.Cells(8, 4).Value, sheet.Cells(8, 5).Value),
            'primary_x': (sheet.Cells(12, 4).Value, sheet.Cells(12, 5).Value),
            'secondary_y': (sheet.Cells(9, 4).Value, sheet.Cells(9, 5).Value)
        }

    try:
        chart.Parent.Activate()

        y_min, y_max = axes_config['primary_y']
        x_min, x_max = axes_config['primary_x']
        sy_min, sy_max = axes_config['secondary_y']

        chart.Axes(2).MinimumScale = y_min
        chart.Axes(2).MaximumScale = y_max
        chart.Axes(1).MinimumScale = x_min
        chart.Axes(1).MaximumScale = x_max

        try:
            chart.Axes(2, 2).MinimumScale = sy_min
            chart.Axes(2, 2).MaximumScale = sy_max
        except Exception:
            print(f"  (No secondary Y axis for Chart{chart_number}, skipping)")

        print(f"Updated axes for Chart{chart_number}")

    except Exception as e:
        print(f"Failed to set axes for Chart{chart_number}: {e}")

def replace_text_in_runs(paragraph, search_text, replace_text):
    for run in paragraph.runs:
        run.text = run.text.replace(search_text, replace_text)

def replace_text_in_document(doc, replacements):
    for para in doc.paragraphs:
        for search_text, replace_text in replacements.items():
            replace_text_in_runs(para, search_text, replace_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for search_text, replace_text in replacements.items():
                        replace_text_in_runs(para, search_text, replace_text)

# -------------------------
# Process Each Device
# -------------------------
excel = win32.Dispatch("Excel.Application")
excel.Visible = False

for _, row in devices_df.iterrows():
    lot_id = str(row["Lot_ID"]).strip()
    dev_num = str(row["Dev#"]).strip()
    # Handle SN: if it's empty or NaN, use dev_num as fallback
    sn_raw = row["SN"]
    if pd.isna(sn_raw) or str(sn_raw).strip() == "":
        sn = dev_num  # Use device number as serial number when SN is empty
    else:
        sn = str(int(sn_raw)).strip()
    sku = str(row["SKU"]).strip()

    print(f"Processing Device: Lot={lot_id}, Dev={dev_num}, SN={sn}, SKU={sku}")

    wb = excel.Workbooks.Open(excel_template_path)
    sheet = wb.Sheets("snl")

    clear_old_data(sheet)

    paste_text_file_fast(sheet, 17, 1, lot_id, dev_num, "WLT_Wave")
    paste_text_file_fast(sheet, 46, 1, lot_id, dev_num, "WLT_SMSR")
    paste_text_file_fast(sheet, 79, 1, lot_id, dev_num, "LIV_vs_Temp")

    sheet.Cells(1, 2).Value = sku

    sheet.Calculate()
    excel.CalculateFull()

    charts_sheet = wb.Sheets("Charts")
    chart1 = charts_sheet.ChartObjects("Chart1").Chart
    chart2 = charts_sheet.ChartObjects("Chart2").Chart

    update_chart_axes(sheet, chart1, 1)
    update_chart_axes(sheet, chart2, 2)

    liv_chart_path = os.path.join(destination_folder, "temp_chart_liv.png")
    smsr_chart_path = os.path.join(destination_folder, "temp_chart_smsr.png")

    chart1.Export(liv_chart_path)
    chart2.Export(smsr_chart_path)

    wb.Close(SaveChanges=False)

    resized_liv_chart_path = liv_chart_path.replace(".png", "_resized.png")
    resized_smsr_chart_path = smsr_chart_path.replace(".png", "_resized.png")
    resize_image(liv_chart_path, resized_liv_chart_path, 130)
    resize_image(smsr_chart_path, resized_smsr_chart_path, 130)

    output_path = os.path.join(data_package_folder, f"{sn} {sku} {dev_num}.docx")
    shutil.copyfile(word_template_path, output_path)

    python_doc = Document(output_path)

    replacements = {"DEV-HERE": dev_num, "SN-HERE": sn, "SKU-HERE": sku, "TODAYS-DATE": datetime.now().strftime("%m/%d/%Y")}
    replace_text_in_document(python_doc, replacements)

    for para in python_doc.paragraphs:
        if "LIV-IMAGE-HERE" in para.text:
            para.text = ""
            para.add_run().add_picture(resized_smsr_chart_path, width=Inches(6))
        elif "SMSR-IMAGE-HERE" in para.text:
            para.text = ""
            para.add_run().add_picture(resized_liv_chart_path, width=Inches(6))

    python_doc.save(output_path)

    os.remove(liv_chart_path)
    os.remove(smsr_chart_path)
    os.remove(resized_liv_chart_path)
    os.remove(resized_smsr_chart_path)

excel.Quit()

print("All datasheets created successfully.")
