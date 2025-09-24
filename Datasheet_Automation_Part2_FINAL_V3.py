import os
import shutil
import pandas as pd

from datetime import datetime
import win32com.client as win32
from docx import Document
from PIL import Image
from docx.shared import Inches

# -------------------------
# CONFIGURATION - Paths
# -------------------------
destination_folder = r"C:\Users\crathod\Documents\Datasheet Automation\Script Output"
data_package_folder = os.path.join(destination_folder, "Data Package")
excel_template_path = r"C:\Users\crathod\Documents\Datasheet Automation\Datasheet Graph Template 1.xlsm"
word_template_path = r"C:\Users\crathod\Documents\Datasheet Automation\Datasheet Template.docx"
os.makedirs(data_package_folder, exist_ok=True)

# -------------------------
# Load Devices.xlsx
# -------------------------
devices_file = os.path.join(destination_folder, "Devices.xlsx")
devices_df = pd.read_excel(devices_file, sheet_name="Devices")
# Don't drop rows with empty SN - we'll handle that in processing
devices_df = devices_df.dropna(subset=["Lot_ID", "Dev#", "SKU"])

# -------------------------
# Helpers
# -------------------------
def find_device_file(lot_id, dev_num, phrase):
    for filename in os.listdir(os.path.join(destination_folder, "Other")):
        if lot_id in filename and dev_num in filename and phrase in filename:
            return os.path.join(destination_folder, "Other", filename)
    return None

def paste_text_file_fast(sheet, start_row, start_column, lot_id, dev_num, phrase):
    file_path = find_device_file(lot_id, dev_num, phrase)
    if not file_path:
        print(f"File not found for {phrase} - skipping.")
        return None

    with open(file_path, 'r') as f:
        lines = [line.strip() for line in f.readlines()]

    data = [line.split() for line in lines]
    num_rows = len(data)
    num_cols = max(len(row) for row in data)

    for row in data:
        while len(row) < num_cols:
            row.append("")

    start_cell = sheet.Cells(start_row, start_column)
    end_cell = sheet.Cells(start_row + num_rows - 1, start_column + num_cols - 1)
    sheet.Range(start_cell, end_cell).Value = data

    print(f"Fast-pasted {len(data)} rows for {phrase} starting at row {start_row}, column {start_column}")
    return file_path

def clear_old_data(sheet):
    """Clear specific rows before pasting new data."""
    print("Clearing old data in rows 69-72 and 40-43...")
    sheet.Range("A69:CB72").ClearContents()
    sheet.Range("A40:CB43").ClearContents()
    print("Old data cleared.\n")

def resize_image(input_path, output_path, scale_percent):
    with Image.open(input_path) as img:
        new_width = int(img.width * (scale_percent / 100))
        new_height = int(img.height * (scale_percent / 100))
        img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        img.save(output_path)

def update_chart_axes(sheet, chart, chart_number):
    if chart_number == 1:
        # Chart 1: Wavelength vs Current (+ SMSR)
        axes_config = {
            'primary_y': (sheet.Cells(13, 4).Value, sheet.Cells(13, 5).Value),  # Wavelength
            'primary_x': (sheet.Cells(7, 4).Value, sheet.Cells(7, 5).Value),    # Current
            'secondary_y': (sheet.Cells(10, 4).Value, sheet.Cells(10, 5).Value) # SMSR
        }
    elif chart_number == 2:
        # Chart 2: LIV Characteristics (Power vs Current + Voltage)
        axes_config = {
            'primary_y': (sheet.Cells(8, 4).Value, sheet.Cells(8, 5).Value),    # Power
            'primary_x': (sheet.Cells(12, 4).Value, sheet.Cells(12, 5).Value),  # Current
            'secondary_y': (sheet.Cells(9, 4).Value, sheet.Cells(9, 5).Value)   # Voltage
        }

    try:
        chart.Parent.Activate()

        y_min, y_max = axes_config['primary_y']
        x_min, x_max = axes_config['primary_x']
        sy_min, sy_max = axes_config['secondary_y']
        
        print(f"Chart {chart_number} - Original values:")
        print(f"  Primary Y: {y_min} to {y_max}")
        print(f"  Secondary Y: {sy_min} to {sy_max}")
        print(f"  X: {x_min} to {x_max}")
        
        # Chart-specific expansions
        if chart_number == 1:
            # Wavelength chart: very small expansion
            sy_min = -sy_max  # SMSR axis should start from negative of max
            y_min = y_min + 0.5 # Expand wavelength min slightly for better view
            y_max_expanded = y_max  # Only +2nm as you specified
            sy_max_expanded = sy_max * 1.2  # SMSR can have more room
            x_max_expanded = x_max * 1.05   # Current: small expansion
            
        else:
            # LIV chart: moderate expansion for power
            y_max_expanded = max(y_max * 1, 0.05)  # Power: ensure at least 0.05W
            sy_max_expanded = sy_max * 1.1  # Voltage: small expansion
            x_max_expanded = x_max * 1.05   # Current: small expansion
        
        # Set axes WITH custom units (full manual control)
        chart.Axes(2).MinimumScale = y_min
        chart.Axes(2).MaximumScale = y_max_expanded
        chart.Axes(1).MinimumScale = x_min
        chart.Axes(1).MaximumScale = x_max_expanded

        # Custom units for Primary Y-axis (Left Y - Wavelength/Power)
        if chart_number == 1:
            # Chart 1: Wavelength axis - use smaller, precise units
            chart.Axes(2).MajorUnit = 0.5  # 0.5nm intervals for wavelength
            chart.Axes(2).MinorUnit = 0.1  # 0.1nm minor intervals
            
            # Ensure minor ticks are visible on wavelength axis
            try:
                chart.Axes(2).MinorTickMark = 2  # xlTickMarkOutside - show minor ticks outside
                chart.Axes(2).HasMinorGridlines = False  # Don't show minor gridlines (cleaner look)
                print(f"  Wavelength axis minor ticks enabled: 0.1nm intervals")
            except Exception as e:
                print(f"  Minor tick configuration failed: {e}")
            
        else:
            # Chart 2: Power axis - use appropriate power units
            chart.Axes(2).MajorUnit = 0.01  # 0.01W intervals for power
            chart.Axes(2).MinorUnit = 0.002  # 0.002W minor intervals

        # Custom units for X-axis (Current) - same for both charts
        chart.Axes(1).MajorUnit = 0.01   # 0.01A intervals for current
        chart.Axes(1).MinorUnit = 0.002  # 0.002A minor intervals

        try:
            chart.Axes(2, 2).MinimumScale = sy_min
            chart.Axes(2, 2).MaximumScale = sy_max_expanded
            
            # Custom units for Secondary Y-axis (Right Y - SMSR/Voltage)
            if chart_number == 1:
                # Chart 1: SMSR axis
                chart.Axes(2, 2).MajorUnit = 10  # 10dB intervals for SMSR
                chart.Axes(2, 2).MinorUnit = 2   # 2dB minor intervals
            else:
                # Chart 2: Voltage axis
                chart.Axes(2, 2).MajorUnit = 0.2  # 0.2V intervals for voltage
                chart.Axes(2, 2).MinorUnit = 0.05  # 0.05V minor intervals
            
            print(f"  Secondary Y major unit: {chart.Axes(2, 2).MajorUnit}")
            
        except Exception as e:
            print(f"  (No secondary Y axis for Chart{chart_number}: {e})")

    except Exception as e:
        print(f"Failed to set axes for Chart{chart_number}: {e}")
        import traceback
        traceback.print_exc()

def replace_text_in_runs(paragraph, search_text, replace_text):
    for run in paragraph.runs:
        run.text = run.text.replace(search_text, replace_text)

def replace_text_in_document(doc, replacements):
    for para in doc.paragraphs:
        for search_text, replace_text in replacements.items():
            replace_text_in_runs(para, search_text, replace_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for search_text, replace_text in replacements.items():
                        replace_text_in_runs(para, search_text, replace_text)

# -------------------------
# Process Each Device
# -------------------------
excel = win32.Dispatch("Excel.Application")
excel.Visible = False

for _, row in devices_df.iterrows():
    lot_id = str(row["Lot_ID"]).strip()
    dev_num = str(row["Dev#"]).strip()
    # Handle SN: if it's empty or NaN, use dev_num as fallback
    sn_raw = row["SN"]
    if pd.isna(sn_raw) or str(sn_raw).strip() == "":
        sn = dev_num  # Use device number as serial number when SN is empty
    else:
        sn = str(int(sn_raw)).strip()
    sku = str(row["SKU"]).strip()

    print(f"Processing Device: Lot={lot_id}, Dev={dev_num}, SN={sn}, SKU={sku}")

    wb = excel.Workbooks.Open(excel_template_path)
    sheet = wb.Sheets("snl")

    clear_old_data(sheet)

    paste_text_file_fast(sheet, 17, 1, lot_id, dev_num, "WLT_Wave")
    paste_text_file_fast(sheet, 46, 1, lot_id, dev_num, "WLT_SMSR")
    paste_text_file_fast(sheet, 79, 1, lot_id, dev_num, "LIV_vs_Temp")

    sheet.Cells(1, 2).Value = sku

    # Control Excel calculation timing
    print("Performing Excel calculations...")
    original_calculation = excel.Calculation
    excel.Calculation = -4135  # xlCalculationManual
    sheet.Calculate()
    excel.CalculateFull()
    
    # Small delay to ensure calculations complete
    import time
    time.sleep(0.5)

    charts_sheet = wb.Sheets("Charts")
    chart1 = charts_sheet.ChartObjects("Chart1").Chart
    chart2 = charts_sheet.ChartObjects("Chart2").Chart

    # Set axes with expanded bounds and units control
    print("Setting chart axes...")
    update_chart_axes(sheet, chart1, 1)
    update_chart_axes(sheet, chart2, 2)
    
    # Verify axes settings before export
    print("Verifying axis settings...")
    try:
        print(f"Chart1 verification - Left Y Max: {chart1.Axes(2).MaximumScale}")
        print(f"Chart1 verification - Right Y Max: {chart1.Axes(2, 2).MaximumScale}")
        print(f"Chart2 verification - Left Y Max: {chart2.Axes(2).MaximumScale}")
    except Exception as e:
        print(f"Verification warning: {e}")
    
    # Small delay before export to ensure settings are applied
    time.sleep(0.5)

    liv_chart_path = os.path.join(destination_folder, "temp_chart_liv.png")
    smsr_chart_path = os.path.join(destination_folder, "temp_chart_smsr.png")

    print("Exporting charts...")
    chart1.Export(liv_chart_path)
    chart2.Export(smsr_chart_path)

    # Restore calculation mode
    excel.Calculation = original_calculation

    wb.Close(SaveChanges=False)

    resized_liv_chart_path = liv_chart_path.replace(".png", "_resized.png")
    resized_smsr_chart_path = smsr_chart_path.replace(".png", "_resized.png")
    resize_image(liv_chart_path, resized_liv_chart_path, 130)
    resize_image(smsr_chart_path, resized_smsr_chart_path, 130)

    output_path = os.path.join(data_package_folder, f"{sn} {sku} {dev_num}.docx")
    print(f"Creating Word document: {output_path}")
    
    try:
        # Copy Word template
        shutil.copyfile(word_template_path, output_path)
        print(f"Template copied successfully")
        
        # Open document
        python_doc = Document(output_path)
        print(f"Document opened successfully")

        # Replace text placeholders
        replacements = {"DEV-HERE": dev_num, "SN-HERE": sn, "SKU-HERE": sku, "TODAYS-DATE": datetime.now().strftime("%m/%d/%Y")}
        replace_text_in_document(python_doc, replacements)
        print(f"Text replacements completed")

        # Insert images
        images_inserted = 0
        for para in python_doc.paragraphs:
            if "LIV-IMAGE-HERE" in para.text:
                para.text = ""
                para.add_run().add_picture(resized_smsr_chart_path, width=Inches(6))
                images_inserted += 1
                print(f"Inserted LIV chart: {resized_smsr_chart_path}")
            elif "SMSR-IMAGE-HERE" in para.text:
                para.text = ""
                para.add_run().add_picture(resized_liv_chart_path, width=Inches(6))
                images_inserted += 1
                print(f"Inserted SMSR chart: {resized_liv_chart_path}")
        
        print(f"Images inserted: {images_inserted}")

        # Save document
        python_doc.save(output_path)
        print(f"Word document saved successfully: {output_path}")
        
        # Verify file exists and has size > 0
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"File exists with size: {file_size} bytes")
        else:
            print("ERROR: File was not created!")
            
    except Exception as e:
        print(f"ERROR creating Word document: {e}")
        import traceback
        traceback.print_exc()

    os.remove(liv_chart_path)
    os.remove(smsr_chart_path)
    os.remove(resized_liv_chart_path)
    os.remove(resized_smsr_chart_path)
    
    print(f"Completed device {dev_num}\n" + "="*50 + "\n")

excel.Quit()

print("All datasheets created successfully.")
